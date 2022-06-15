"""Добавление изображений"""
import docx

doc = docx.Document()

doc.add_paragraph('Классная картинка снизу')

doc.add_picture('pigeons.png', width=docx.shared.Cm(10))

doc.save('example4.docx')