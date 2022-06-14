"""Создание Word документов"""
import docx
from docx.shared import Pt
doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(14)
doc.add_paragraph('Абзац!', 'List Bullet')

par1 = doc.add_paragraph('Первый абзац.')
par2 = doc.add_paragraph('Второй абзац.')
par3 = doc.add_paragraph('Третий абзац.')

par1.add_run(' Дополнение первого абзаца.').italic = True
par2.add_run(' Дополнение второго абзаца.').bold = True
run = par3.add_run(' Дополнение третьего абзаца.')
run.bold = True
run.underline = True




doc.save('example.docx')