"""Чтение из Word документов"""
import docx
from docx.shared import Pt
doc = docx.Document('example.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[2].text)

print(doc.paragraphs[1].runs[1].text)
print('--------------------')
for paragrath in doc.paragraphs:
    print(paragrath.text)