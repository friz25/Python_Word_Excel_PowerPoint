"""Добавление ТАБЛИЦ """
import docx

doc = docx.Document()

table = doc.add_table(rows=5, cols=3)

table.style = 'Table Grid'

for row in range(5):
    for col in range(3):
        cell = table.cell(row, col)
        cell.text = str(row)

doc.save('example5.docx')
"""Выводим таблицу из WORD файла """
doc = docx.Document('example5.docx')

table = doc.tables[0]

for row in table.rows:
    string = ''
    for cell in row.cells:
        string += cell.text + ' '
    print(string)

