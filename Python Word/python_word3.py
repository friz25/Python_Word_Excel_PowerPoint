"""Добавление заголовков, разрывов страницы
"""
import docx

doc = docx.Document()

# doc.add_heading('Заголовок 0', 0)
# doc.add_heading('Заголовок 1', 1)
# doc.add_heading('Заголовок 2', 2)
# doc.add_heading('Заголовок 3', 3)
# doc.add_heading('Заголовок 4', 4)

doc.add_paragraph('Это первая страница')
doc.add_page_break()
doc.add_paragraph('Это вторая страница')

doc.save('example.docx')